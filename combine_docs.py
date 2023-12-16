from docx.text.paragraph import Paragraph
from docx import Document
from docx.table import Table
from copy import deepcopy


def iter_headings(blocks):
    """
    Iterate over headings
    """
    idx = 0
    for block in blocks:
        if isinstance(block, Paragraph):
            if block.style.name.startswith('Heading'):
                yield block, idx
        idx += 1

def delete_paragraph(paragraph):
    """ delete paragraph"""
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def delete_table(document, table):
    """ delete table"""
    table._element.getparent().remove(table._element)

def move_table_after(table, paragraph):
    """ move table after a paragraph """
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

def get_header_indices(document, heading_name=None):
    id_list = {}
    case_found = False
    all_blocks = [x for x in document.iter_inner_content()] 
    for heading, idx in iter_headings(all_blocks):
        if case_found is True:
            case_found = False
            id_list[last_txt].append(idx - 1)
            id_list[last_txt] = tuple(id_list[last_txt])
        if heading_name in heading.text:
            id_list[heading_name] = [idx]
            last_txt = heading_name
            case_found = True
    return id_list

# Open the two Word documents that you want to compare and combine.
doc1 = Document("EE207 Assignment 1.docx")
doc2 = Document("EE207 Assignment 2.docx")
headings_list = ["To amend default styles:"]
print(len(doc1.paragraphs))
print(len(doc2.paragraphs))
for heading in headings_list:
    ids_doc1 = get_header_indices(doc1, heading_name=heading) 
    ids_doc2 = get_header_indices(doc2, heading_name=heading)
    all_blocks1 =[x for x in doc1.iter_inner_content()]
    for item, indices in ids_doc1.items():
        print(ids_doc1)
        for idx in range(indices[1], indices[0] - 1, -1):
            if isinstance(all_blocks1[idx], Paragraph):
                delete_paragraph(all_blocks1[idx])
            else:
                delete_table(doc1, all_blocks1[idx])
        prev_idx = indices[0] - 1
    
    for i in range(prev_idx, -1, -1):
        if isinstance(all_blocks1[i], Paragraph): 
            last_para = all_blocks1[i]
            break

    print(len(doc1.paragraphs))
    print(len(doc2.paragraphs))
    for item, indices in ids_doc2.items():
        print(ids_doc2)
        all_blocks2 = [x for x in doc2.iter_inner_content()]
        for idx in range(indices[0], indices[1] + 1, 1):
            print(type(all_blocks2[idx]))
            if isinstance(all_blocks2[idx], Paragraph):
                para_doc1 = doc1.add_paragraph()
                pic_found = False 
                for run in all_blocks2[idx].runs:
                    pic_found_in_run = False 
                    for inline in run._r.xpath("w:drawing/wp:inline"):
                        width = float(inline.extent.cx) # in EMUs https://startbigthinksmall.wordpress.com/2010/01/04/points-inches-and-emus-measuring-units-in-office-open-xml/
                        height = float(inline.extent.cy)
                        rId = inline.graphic.graphicData.pic.blipFill.blip.embed
                        image = doc2.part.related_parts[rId].image
                        filename = image.filename 
                        with open(filename, 'wb') as f:  # make a copy in the local dir
                            f.write(image.blob)
                        print(', '.join([
                            f"saved image {filename}",
                            f"type {image.content_type}",
                            f"px: {image.px_height} x {image.px_width}",
                            f"size in document: {height} x {width}",
                        ]))
                        para_doc1.add_run().add_picture(filename,  width, height)
                        pic_found = True 
                        pic_found_in_run = True
                    if pic_found_in_run is False:
                        new_run = para_doc1.add_run()
                        new_run._r = deepcopy(run._r)
                if pic_found is False:
                    para_doc1._p = deepcopy(all_blocks2[idx]._p)
                last_para._p.addnext(para_doc1._p)
                last_para = para_doc1
            else:
                print("Added a table")
                move_table_after(deepcopy(all_blocks2[idx]), last_para)
                blank_para = doc1.add_paragraph()
                last_para._p.addnext(blank_para._p)
                last_para = blank_para

print(len(doc1.paragraphs))
print(len(doc2.paragraphs))
doc1.save("New1.docx")
doc2.save("New2.docx")
#doc_new.save("New3.docx")

